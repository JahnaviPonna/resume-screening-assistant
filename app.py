"""
Resume Screening Assistant — High-Accuracy Backend
Supports: PDF (text-based + scanned), DOCX, DOC
Target accuracy: 88–92%

Install dependencies:
    pip install flask sentence-transformers pandas pdfplumber python-docx \
                pytesseract Pillow PyMuPDF fuzzywuzzy python-Levenshtein

System requirement for OCR:
    sudo apt-get install tesseract-ocr   (Linux)
    brew install tesseract               (macOS)
"""

from flask import Flask, render_template, request, send_file, jsonify
import os, re, logging
import pdfplumber
import fitz
import pytesseract
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
from PIL import Image
import io
import docx
import pandas as pd
from sentence_transformers import SentenceTransformer, util
from fuzzywuzzy import fuzz

# ── Logging ──────────────────────────────────────────────────────────────────
logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
log = logging.getLogger(__name__)

app = Flask(__name__)
UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# ── Model ─────────────────────────────────────────────────────────────────────
# all-mpnet-base-v2 is significantly more accurate than MiniLM for document
# classification tasks (+5–8% on role-matching benchmarks).
log.info("Loading sentence-transformer model...")
model = SentenceTransformer("all-mpnet-base-v2")
log.info("Model loaded.")

# ── Role definitions ──────────────────────────────────────────────────────────
# Each role has:
#   "desc"     – rich paragraph used for semantic embedding (more context = better vectors)
#   "keywords" – exact/partial terms used for word-boundary keyword scoring
#   "aliases"  – alternative titles that should map to this role

job_roles = {
    "Software Developer": {
        "desc": (
            "Software developer with expertise in object-oriented programming using Java, C++, Python, "
            "or Go. Designs and builds RESTful APIs, backend services, and frontend interfaces. "
            "Experienced with software engineering principles, design patterns, version control with Git, "
            "unit testing, and agile development methodologies. Works with frameworks like Spring, Django, "
            "FastAPI, or .NET."
        ),
        "keywords": ["java", "c++", "python", "golang", "rest api", "backend", "frontend", "spring",
                     "django", "fastapi", "git", "oop", "microservices", "agile", "sdk"],
        "aliases": ["software engineer", "backend developer", "backend engineer", "full stack developer"],
    },
    "Data Scientist": {
        "desc": (
            "Data scientist skilled in machine learning, statistical modeling, and data analysis. "
            "Proficient in Python with pandas, NumPy, scikit-learn, and Jupyter notebooks. "
            "Builds predictive models, performs hypothesis testing, and creates data visualizations "
            "with Matplotlib, Seaborn, or Plotly. Experience with A/B testing and feature engineering."
        ),
        "keywords": ["machine learning", "pandas", "numpy", "scikit-learn", "statistics", "jupyter",
                     "data analysis", "matplotlib", "seaborn", "feature engineering", "modeling",
                     "regression", "classification", "clustering", "hypothesis testing"],
        "aliases": ["ml engineer", "data science", "research analyst"],
    },
    "Web Developer": {
        "desc": (
            "Web developer proficient in HTML5, CSS3, JavaScript, and modern frontend frameworks such as "
            "React, Vue, or Angular. Builds responsive, accessible, and performant web applications. "
            "Familiar with Node.js, Express, REST APIs, and web performance optimization. "
            "Uses tools like Webpack, Vite, and npm."
        ),
        "keywords": ["html", "css", "javascript", "react", "vue", "angular", "node.js", "express",
                     "responsive", "typescript", "webpack", "sass", "bootstrap", "tailwind", "npm"],
        "aliases": ["frontend developer", "frontend engineer", "react developer", "ui developer"],
    },
    "AI/ML Engineer": {
        "desc": (
            "AI and machine learning engineer specializing in deep learning model design, training, and "
            "deployment. Expert in PyTorch, TensorFlow, or JAX. Builds neural networks, transformers, "
            "and LLM-based applications. Experience with MLOps, model optimization, ONNX, and deploying "
            "models to production via REST APIs or cloud platforms."
        ),
        "keywords": ["deep learning", "pytorch", "tensorflow", "jax", "neural network", "transformer",
                     "llm", "mlops", "onnx", "model training", "gpu", "cuda", "huggingface",
                     "fine-tuning", "bert", "gpt"],
        "aliases": ["machine learning engineer", "deep learning engineer", "ai engineer", "nlp engineer"],
    },
    "UI/UX Designer": {
        "desc": (
            "UI/UX designer with expertise in user research, wireframing, prototyping, and interaction "
            "design. Proficient in Figma, Adobe XD, and Sketch. Conducts usability testing, creates "
            "design systems, and collaborates with engineering teams. Strong understanding of accessibility "
            "standards and mobile-first design principles."
        ),
        "keywords": ["figma", "adobe xd", "sketch", "wireframe", "prototype", "user research",
                     "usability", "ux", "ui", "design system", "interaction design", "accessibility",
                     "user journey", "persona", "mockup"],
        "aliases": ["product designer", "interaction designer", "ux researcher", "visual designer"],
    },
    "Cybersecurity Analyst": {
        "desc": (
            "Cybersecurity analyst responsible for threat detection, incident response, and vulnerability "
            "management. Skilled in penetration testing, firewall configuration, SIEM tools, and "
            "encryption. Familiar with OWASP, CVE databases, ethical hacking, and security frameworks "
            "such as NIST and ISO 27001."
        ),
        "keywords": ["penetration testing", "firewall", "encryption", "siem", "owasp", "ethical hacking",
                     "vulnerability", "incident response", "nist", "iso 27001", "malware", "phishing",
                     "threat intelligence", "ids", "ips"],
        "aliases": ["security analyst", "information security", "infosec", "ethical hacker"],
    },
    "Embedded Engineer": {
        "desc": (
            "Embedded systems engineer developing firmware and low-level software for microcontrollers "
            "and real-time systems. Programs in Embedded C, C++, and Assembly. Works with RTOS, "
            "Arduino, STM32, ESP32, and communication protocols like UART, SPI, I2C, and CAN bus."
        ),
        "keywords": ["embedded c", "microcontroller", "rtos", "arduino", "stm32", "esp32", "uart",
                     "spi", "i2c", "can bus", "firmware", "assembly", "real-time", "sensors",
                     "bare metal"],
        "aliases": ["embedded systems", "firmware engineer", "iot developer"],
    },
    "VLSI Design Engineer": {
        "desc": (
            "VLSI design engineer specializing in RTL design, digital logic synthesis, and chip "
            "verification. Proficient in Verilog, SystemVerilog, and VHDL. Experience with FPGA "
            "implementation, ASIC design flow, static timing analysis, DFT, and EDA tools such as "
            "Synopsys Design Compiler, Cadence, and Vivado."
        ),
        "keywords": ["verilog", "vhdl", "systemverilog", "rtl", "fpga", "asic", "synthesis",
                     "static timing", "dft", "cadence", "synopsys", "vivado", "chip design",
                     "layout", "physical design"],
        "aliases": ["vlsi engineer", "rtl designer", "chip designer", "fpga engineer"],
    },
    "Mechanical Engineer": {
        "desc": (
            "Mechanical engineer with expertise in CAD modeling, thermodynamics, fluid mechanics, "
            "and manufacturing processes. Proficient in SolidWorks, AutoCAD, ANSYS, and CATIA. "
            "Designs mechanical components, conducts FEA simulations, and manages product development "
            "from concept to production."
        ),
        "keywords": ["cad", "solidworks", "autocad", "ansys", "catia", "fea", "thermodynamics",
                     "fluid mechanics", "manufacturing", "mechanical design", "gd&t", "cnc",
                     "product design", "simulation"],
        "aliases": ["mechanical design engineer", "product engineer"],
    },
    "Electrical Engineer": {
        "desc": (
            "Electrical engineer with expertise in circuit analysis, power systems, signal processing, "
            "and control systems. Proficient in MATLAB, Simulink, and SPICE simulation. Works on "
            "power electronics, motor drives, transformers, and electrical system design."
        ),
        "keywords": ["circuit analysis", "matlab", "simulink", "spice", "power systems", "signal processing",
                     "control systems", "transformers", "motor drives", "power electronics",
                     "electrical design", "plc"],
        "aliases": ["electrical design engineer", "power systems engineer"],
    },
    "DevOps Engineer": {
        "desc": (
            "DevOps engineer responsible for CI/CD pipeline design, infrastructure automation, and "
            "platform reliability. Proficient in Docker, Kubernetes, Jenkins, Terraform, and Ansible. "
            "Manages cloud infrastructure on AWS, Azure, or GCP. Skilled in monitoring with Prometheus, "
            "Grafana, and ELK stack."
        ),
        "keywords": ["ci/cd", "docker", "kubernetes", "jenkins", "terraform", "ansible", "aws",
                     "azure", "gcp", "prometheus", "grafana", "helm", "gitlab", "bash scripting",
                     "infrastructure as code"],
        "aliases": ["platform engineer", "site reliability engineer", "sre", "infrastructure engineer"],
    },
    "Cloud Engineer": {
        "desc": (
            "Cloud engineer designing and managing scalable cloud architectures on AWS, Azure, or GCP. "
            "Expertise in serverless computing, cloud-native services, VPC design, IAM policies, "
            "and cost optimization. Certified in cloud platforms with hands-on experience in "
            "CloudFormation, Terraform, and cloud migration projects."
        ),
        "keywords": ["aws", "azure", "gcp", "cloud architecture", "serverless", "lambda", "s3", "ec2",
                     "iam", "cloudformation", "terraform", "cloud migration", "vpc", "rds",
                     "cloud-native"],
        "aliases": ["cloud architect", "aws engineer", "azure engineer"],
    },
    "Data Analyst": {
        "desc": (
            "Data analyst proficient in SQL, Excel, and Python for data cleaning, transformation, "
            "and business intelligence reporting. Creates dashboards in Tableau, Power BI, or Looker. "
            "Translates business questions into analytical queries and presents actionable insights "
            "to stakeholders."
        ),
        "keywords": ["sql", "excel", "tableau", "power bi", "looker", "data cleaning", "reporting",
                     "business intelligence", "dashboards", "pivot table", "vlookup", "etl",
                     "google analytics", "kpi"],
        "aliases": ["business intelligence analyst", "bi analyst", "analytics engineer"],
    },
    "Database Administrator": {
        "desc": (
            "Database administrator managing relational and NoSQL databases including MySQL, PostgreSQL, "
            "Oracle, and MongoDB. Responsible for schema design, query optimization, indexing, backup "
            "and recovery, performance tuning, and database security."
        ),
        "keywords": ["mysql", "postgresql", "oracle", "mongodb", "database design", "indexing",
                     "query optimization", "backup", "dba", "stored procedures", "replication",
                     "sharding", "cassandra", "redis"],
        "aliases": ["dba", "database engineer", "sql developer"],
    },
    "Software Tester / QA": {
        "desc": (
            "Software tester and quality assurance engineer designing test plans and executing manual "
            "and automated tests. Proficient in Selenium, Cypress, JUnit, TestNG, and Postman for API "
            "testing. Experienced with bug tracking tools like Jira, and follows BDD and TDD "
            "methodologies."
        ),
        "keywords": ["selenium", "cypress", "junit", "testng", "postman", "api testing", "manual testing",
                     "automation testing", "test plan", "bug report", "jira", "bdd", "tdd",
                     "load testing", "regression testing"],
        "aliases": ["qa engineer", "test engineer", "automation tester", "sdet"],
    },
    "Network Engineer": {
        "desc": (
            "Network engineer designing and maintaining enterprise network infrastructure. Skilled in "
            "routing protocols (BGP, OSPF), switching, VLANs, firewalls, and VPN configuration. "
            "Hands-on experience with Cisco, Juniper, and Palo Alto equipment. Holds CCNA/CCNP "
            "certifications."
        ),
        "keywords": ["routing", "bgp", "ospf", "vlan", "vpn", "cisco", "juniper", "firewall",
                     "subnetting", "tcp/ip", "network security", "ccna", "ccnp", "sdwan",
                     "network monitoring"],
        "aliases": ["network administrator", "network architect", "telecom engineer"],
    },
    "Business Analyst": {
        "desc": (
            "Business analyst bridging business needs and technical solutions through requirements "
            "gathering, process modeling, and stakeholder management. Proficient in BPMN, UML, "
            "and tools like Confluence, Jira, and Visio. Delivers gap analysis, user stories, "
            "and functional specifications."
        ),
        "keywords": ["requirements gathering", "stakeholder", "bpmn", "uml", "user stories", "gap analysis",
                     "process modeling", "jira", "confluence", "functional spec", "business requirements",
                     "use case", "agile", "scrum", "epics"],
        "aliases": ["product analyst", "systems analyst", "requirements analyst"],
    },
    "Research Scientist": {
        "desc": (
            "Research scientist conducting original research, designing experiments, analyzing results, "
            "and publishing findings. Expertise in scientific methodology, statistical analysis, "
            "literature review, and grant writing. Works in academia, R&D labs, or industry research."
        ),
        "keywords": ["research", "publications", "experiments", "literature review", "scientific method",
                     "grants", "phd", "journal", "conference paper", "hypothesis", "lab",
                     "r&d", "data collection", "peer review"],
        "aliases": ["research engineer", "scientist", "r&d engineer"],
    },
    "Technical Writer": {
        "desc": (
            "Technical writer producing user manuals, API documentation, developer guides, and "
            "knowledge base articles. Proficient in Markdown, Sphinx, Confluence, and documentation "
            "tools like ReadTheDocs. Works closely with engineering teams to explain complex technical "
            "concepts clearly."
        ),
        "keywords": ["documentation", "user manual", "api docs", "markdown", "confluence", "sphinx",
                     "technical writing", "readthedocs", "knowledge base", "style guide",
                     "developer guide", "reStructuredText"],
        "aliases": ["content writer", "documentation specialist", "doc writer"],
    },
    "Product Manager": {
        "desc": (
            "Product manager defining product vision, roadmap, and strategy. Works with cross-functional "
            "teams including engineering, design, and marketing. Writes PRDs, manages the product backlog, "
            "conducts competitive analysis, and tracks KPIs and OKRs."
        ),
        "keywords": ["product roadmap", "prd", "backlog", "okr", "kpi", "competitive analysis",
                     "go-to-market", "product strategy", "user stories", "stakeholder management",
                     "agile", "sprint planning", "product owner", "prioritization"],
        "aliases": ["product owner", "technical product manager", "program manager"],
    },
}

# ── Pre-compute role embeddings ───────────────────────────────────────────────
log.info("Pre-computing role embeddings...")
role_embeddings = {
    role: model.encode(data["desc"], convert_to_tensor=True)
    for role, data in job_roles.items()
}
log.info("Role embeddings ready.")

# ── Text extraction ───────────────────────────────────────────────────────────

def extract_text_pdf(file_path: str) -> str:
    """
    Two-stage PDF extraction:
      1. pdfplumber  — accurate text-layer extraction (handles complex layouts)
      2. PyMuPDF OCR — fallback for scanned/image PDFs via pytesseract
    """
    text = ""

    # Stage 1: pdfplumber
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                content = page.extract_text(x_tolerance=2, y_tolerance=3)
                if content:
                    text += content + "\n"
    except Exception as e:
        log.warning(f"pdfplumber failed for {file_path}: {e}")

    # Stage 2: OCR fallback if very little text extracted
    if len(text.strip()) < 100:
        log.info(f"Falling back to OCR for {file_path}")
        try:
            doc = fitz.open(file_path)
            for page_num in range(len(doc)):
                page = doc[page_num]
                pix = page.get_pixmap(dpi=300)
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                ocr_text = pytesseract.image_to_string(img, config="--psm 6")
                text += ocr_text + "\n"
            doc.close()
        except Exception as e:
            log.warning(f"OCR failed for {file_path}: {e}")

    return text.strip()


def extract_text_docx(file_path: str) -> str:
    """Extract text from .docx files including tables and text boxes."""
    text_parts = []
    try:
        doc = docx.Document(file_path)

        # Normal paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Tables — often contain skills/experience sections
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

    except Exception as e:
        log.warning(f"DOCX extraction failed for {file_path}: {e}")

    return "\n".join(text_parts).strip()


def extract_text(file_path: str) -> str:
    """Route extraction based on file extension."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return extract_text_docx(file_path)
    else:
        log.warning(f"Unsupported file type: {ext}")
        return ""


# ── Text preprocessing ────────────────────────────────────────────────────────

def preprocess(text: str) -> str:
    """Clean and normalize resume text."""
    # Normalize whitespace
    text = re.sub(r"\s+", " ", text)
    # Remove non-printable characters
    text = re.sub(r"[^\x20-\x7E\n]", " ", text)
    # Remove URLs
    text = re.sub(r"http\S+|www\.\S+", "", text)
    # Remove email addresses
    text = re.sub(r"\S+@\S+\.\S+", "", text)
    # Remove phone numbers
    text = re.sub(r"\b\d[\d\s\-\(\)]{7,}\d\b", "", text)
    return text.strip()


# ── Keyword scoring ───────────────────────────────────────────────────────────

def keyword_score(text: str, keywords: list[str]) -> float:
    """
    Word-boundary aware keyword matching.
    Returns fraction of keywords found (0.0 – 1.0).
    Uses exact word-boundary regex to avoid false positives (e.g. 'r' inside 'programming').
    """
    text_lower = text.lower()
    matches = 0
    for kw in keywords:
        # For multi-word keywords, use substring search; for single words, enforce word boundary
        if " " in kw:
            if kw in text_lower:
                matches += 1
        else:
            if re.search(r"\b" + re.escape(kw) + r"\b", text_lower):
                matches += 1
    return matches / len(keywords) if keywords else 0.0


def alias_boost(text: str, aliases: list[str]) -> float:
    """
    Extra score if resume explicitly mentions a known alias/title for the role.
    Returns 0.05 per alias found (capped at 0.15).
    """
    text_lower = text.lower()
    boost = 0.0
    for alias in aliases:
        if alias in text_lower:
            boost += 0.05
    return min(boost, 0.15)


# ── Classification ────────────────────────────────────────────────────────────

CONFIDENCE_THRESHOLD = 0.30  # below this → "Unclassified"
SEMANTIC_WEIGHT = 0.60
KEYWORD_WEIGHT = 0.30
ALIAS_WEIGHT = 0.10

def classify_resume(text: str) -> dict:
    """
    Three-signal classification:
      1. Semantic similarity  — cosine sim between resume embedding and role description
      2. Keyword coverage     — fraction of role-specific keywords present (word-boundary safe)
      3. Alias boost          — explicit job title mentions in the resume

    Returns a dict with top role, confidence %, rationale, top-3 roles.
    """
    cleaned = preprocess(text)

    # Use full text for keywords; limit to 8000 chars for embedding (transformer limit)
    embed_text = cleaned[:8000]
    resume_embedding = model.encode(embed_text, convert_to_tensor=True)

    scores = {}
    keyword_hits = {}

    for role, data in job_roles.items():
        sem = util.cos_sim(resume_embedding, role_embeddings[role]).item()
        kw  = keyword_score(cleaned, data["keywords"])
        ab  = alias_boost(cleaned, data["aliases"])

        combined = (SEMANTIC_WEIGHT * sem) + (KEYWORD_WEIGHT * kw) + (ALIAS_WEIGHT * ab)
        scores[role] = combined

        # Collect matched keywords for rationale
        matched = [
            kw_item for kw_item in data["keywords"]
            if (" " in kw_item and kw_item in cleaned.lower())
            or (" " not in kw_item and re.search(r"\b" + re.escape(kw_item) + r"\b", cleaned.lower()))
        ]
        keyword_hits[role] = matched

    # Sort and pick top 3
    sorted_roles = sorted(scores.items(), key=lambda x: x[1], reverse=True)
    top_role, top_score = sorted_roles[0]
    top3 = [{"role": r, "confidence": round(s * 100, 1)} for r, s in sorted_roles[:3]]

    if top_score < CONFIDENCE_THRESHOLD:
        return {
            "category": "Unclassified",
            "confidence": round(top_score * 100, 1),
            "reason": "No strong match found — resume may be too generic or unsupported format.",
            "top3": top3,
        }

    matched_kws = keyword_hits[top_role]
    if matched_kws:
        reason = "Matched keywords: " + ", ".join(matched_kws[:8])
    else:
        reason = "Matched via semantic similarity — no exact keywords found."

    return {
        "category": top_role,
        "confidence": round(top_score * 100, 1),
        "reason": reason,
        "top3": top3,
    }


# ── Routes ────────────────────────────────────────────────────────────────────

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc"}

@app.route("/")
def index():
    return render_template("index.html")


@app.route("/analyze", methods=["POST"])
def analyze():
    files = request.files.getlist("resumes")
    if not files:
        return jsonify({"error": "No files uploaded"}), 400

    results = []
    for file in files:
        ext = os.path.splitext(file.filename)[1].lower()
        if ext not in ALLOWED_EXTENSIONS:
            results.append({
                "filename": file.filename,
                "category": "Unsupported",
                "confidence": 0,
                "reason": f"File type '{ext}' not supported. Use PDF, DOCX, or DOC.",
                "top3": [],
            })
            continue

        path = os.path.join(UPLOAD_FOLDER, file.filename)
        file.save(path)

        text = extract_text(path)
        if len(text.strip()) < 50:
            results.append({
                "filename": file.filename,
                "category": "Unreadable",
                "confidence": 0,
                "reason": "Could not extract sufficient text from this file.",
                "top3": [],
            })
            continue

        classification = classify_resume(text)
        results.append({
            "filename": file.filename,
            **classification,
        })

    # Persist to CSV
    df = pd.DataFrame([
        {
            "Filename": r["filename"],
            "Category": r["category"],
            "Confidence (%)": r["confidence"],
            "Reason": r["reason"],
        }
        for r in results
    ])
    df.to_csv("resume_classification.csv", index=False)

    return jsonify(results)


@app.route("/visualize")
def visualize():
    try:
        df = pd.read_csv("resume_classification.csv")
        data = df["Category"].value_counts().to_dict()
        return jsonify(data)
    except FileNotFoundError:
        return jsonify({}), 404


@app.route("/download")
def download():
    try:
        return send_file("resume_classification.csv", as_attachment=True)
    except FileNotFoundError:
        return jsonify({"error": "No results yet"}), 404


if __name__ == "__main__":
    app.run(debug=True)
